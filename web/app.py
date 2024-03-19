import eel
from docx import Document

dom = Document('./files/tamplate.docx')
list = []
for text in dom.paragraphs:
    list.append(text.text)

@eel.expose
def run(input):

    for i in range(len(input)):
        for j in range(len(input[i])):
           list[i] = nth_repl(list[i], 'X', input[i][j], 1)

    doc = Document()
    for i in list:
        doc.add_paragraph(i)
    doc.save('./files/demo.docx')

def nth_repl(s, sub, repl, n):
    find = s.find(sub)
    i = find != -1
    while find != -1 and i != n:
        find = s.find(sub, find + 1)
        i += 1
    if i == n:
        return s[:find] + repl + s[find+len(sub):]
    return s    

eel.init('web')
eel.start('index.html')
